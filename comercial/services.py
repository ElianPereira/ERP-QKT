from decimal import Decimal
import math
from django.conf import settings
from .models import ItemCotizacion, ConstanteSistema

class CalculadoraBarraService:
    """
    Servicio encargado de toda la lógica de cálculo de barra, 
    separando la lógica de negocio del modelo de base de datos.
    """

    def __init__(self, cotizacion):
        self.cot = cotizacion

    def _get_costo(self, insumo, clave_constante, default_val):
        if insumo:
            factor = insumo.factor_rendimiento if insumo.factor_rendimiento > 0 else 1
            return insumo.costo_unitario / Decimal(factor)
        try:
            const = ConstanteSistema.objects.get(clave=clave_constante)
            return const.valor
        except ConstanteSistema.DoesNotExist:
            return Decimal(default_val)

    def calcular(self):
        c = self.cot
        checks = {
            'refrescos': c.incluye_refrescos,
            'cerveza': c.incluye_cerveza,
            'nacional': c.incluye_licor_nacional,
            'premium': c.incluye_licor_premium,
            'coctel_base': c.incluye_cocteleria_basica,
            'coctel_prem': c.incluye_cocteleria_premium
        }

        if not any(checks.values()) or c.num_personas <= 0:
            return None

        C_HIELO = self._get_costo(c.insumo_hielo, 'PRECIO_HIELO_20KG', '90.00')
        C_MIXER = self._get_costo(c.insumo_refresco, 'PRECIO_REFRESCO_2L', '22.00')
        C_AGUA = self._get_costo(c.insumo_agua, 'PRECIO_AGUA_GAL', '10.00')
        C_ALC_NAC = self._get_costo(c.insumo_alcohol_basico, 'PRECIO_ALC_NAC', '380.00')
        C_ALC_PREM = self._get_costo(c.insumo_alcohol_premium, 'PRECIO_ALC_PREM', '1150.00')
        C_CERVEZA = Decimal('42.00')
        C_GIN = Decimal('550.00')
        C_INSUMO_COCTEL_BASE = Decimal('15.00')
        C_INSUMO_COCTEL_PREM = Decimal('28.00')
        C_EXTRA_BARRA = self._get_costo(None, 'COSTO_EXTRA_BARRA', '0.00')

        R_BOTELLA = 16.0
        R_CAGUAMA = 3.0

        factor_termico = 1.0
        tragos_ph = 1.3

        if c.clima == 'calor':
            factor_termico = 1.3
            tragos_ph = 1.5
        elif c.clima == 'extremo':
            factor_termico = 1.6
            tragos_ph = 1.8

        TOTAL_TRAGOS = c.num_personas * c.horas_servicio * tragos_ph

        pesos = {}
        if checks['cerveza']: pesos['cerveza'] = 55
        if checks['nacional']: pesos['nacional'] = 35
        if checks['premium']: pesos['premium'] = 25
        if checks['coctel_base']: pesos['coctel_base'] = 20
        if checks['coctel_prem']: pesos['coctel_prem'] = 15
        if checks['refrescos']:
            if not pesos: pesos['refrescos'] = 100
            else: pesos['refrescos'] = 15

        total_peso = sum(pesos.values()) or 1

        res = {
            'botellas_nacional': 0, 'botellas_premium': 0, 'cervezas_unidades': 0,
            'litros_mezcladores': 0, 'hielo_consumo_kg': 0.0, 'hielo_enfriamiento_kg': 0.0,
            'bolsas_hielo_20kg': 0, 'costo_alcohol': Decimal(0), 'costo_insumos_varios': Decimal(0)
        }

        costo_puro = {k: Decimal(0) for k in ['cerveza','nacional','premium','coctel','refrescos']}
        litros_mixer_calc = 0.0
        costo_fruta = Decimal(0)

        if 'cerveza' in pesos:
            share = pesos['cerveza'] / total_peso
            tragos = TOTAL_TRAGOS * share
            res['cervezas_unidades'] = math.ceil(tragos / R_CAGUAMA)
            costo = res['cervezas_unidades'] * C_CERVEZA
            res['costo_alcohol'] += costo
            costo_puro['cerveza'] += costo

        if 'nacional' in pesos:
            share = pesos['nacional'] / total_peso
            tragos = TOTAL_TRAGOS * share
            botellas = math.ceil(tragos / R_BOTELLA)
            res['botellas_nacional'] += botellas
            costo = botellas * C_ALC_NAC
            res['costo_alcohol'] += costo
            costo_puro['nacional'] += costo
            litros_mixer_calc += (tragos * 0.200)
            costo_puro['nacional'] += (Decimal(tragos * 0.200) * C_MIXER)

        if 'premium' in pesos:
            share = pesos['premium'] / total_peso
            tragos = TOTAL_TRAGOS * share
            botellas = math.ceil(tragos / R_BOTELLA)
            res['botellas_premium'] += botellas
            costo = botellas * C_ALC_PREM
            res['costo_alcohol'] += costo
            costo_puro['premium'] += costo
            litros_mixer_calc += (tragos * 0.180)
            costo_puro['premium'] += (Decimal(tragos * 0.180) * C_MIXER)

        if 'coctel_base' in pesos:
            share = pesos['coctel_base'] / total_peso
            tragos = TOTAL_TRAGOS * share
            c_ins = Decimal(tragos) * C_INSUMO_COCTEL_BASE
            costo_fruta += c_ins
            costo_puro['coctel'] += c_ins
            b_ron_teq = math.ceil(tragos / R_BOTELLA)
            res['botellas_nacional'] += b_ron_teq
            c_alc = b_ron_teq * C_ALC_NAC
            res['costo_alcohol'] += c_alc
            costo_puro['coctel'] += c_alc
            litros_mixer_calc += (tragos * 0.100)

        if 'coctel_prem' in pesos:
            share = pesos['coctel_prem'] / total_peso
            tragos = TOTAL_TRAGOS * share
            c_ins = Decimal(tragos) * C_INSUMO_COCTEL_PREM
            costo_fruta += c_ins
            costo_puro['coctel'] += c_ins
            b_gin = math.ceil((tragos * 0.3) / R_BOTELLA)
            res['botellas_premium'] += b_gin
            c_alc = b_gin * C_GIN
            res['costo_alcohol'] += c_alc
            costo_puro['coctel'] += c_alc

        if 'refrescos' in pesos:
            share = pesos['refrescos'] / total_peso
            tragos = TOTAL_TRAGOS * share
            litros_mixer_calc += (tragos * 0.355)

        res['litros_mezcladores'] = math.ceil(litros_mixer_calc)
        c_mixers_total = res['litros_mezcladores'] * C_MIXER

        hielo_consumo = TOTAL_TRAGOS * 0.25
        hielo_enfriamiento = 0.0
        if res['cervezas_unidades'] > 0:
            hielo_enfriamiento += (res['cervezas_unidades'] / 30.0) * 20.0
        volumen_a_enfriar = res['litros_mezcladores'] + (c.num_personas * 0.6)
        hielo_enfriamiento += (volumen_a_enfriar / 60.0) * 20.0

        res['hielo_consumo_kg'] = hielo_consumo * factor_termico
        res['hielo_enfriamiento_kg'] = hielo_enfriamiento * factor_termico
        total_hielo_kg = res['hielo_consumo_kg'] + res['hielo_enfriamiento_kg']
        res['bolsas_hielo_20kg'] = math.ceil(total_hielo_kg / 20.0)
        costo_hielo = res['bolsas_hielo_20kg'] * C_HIELO

        litros_agua = math.ceil(c.num_personas * 0.6)
        res['litros_agua'] = litros_agua
        costo_agua = litros_agua * C_AGUA

        res['costo_insumos_varios'] = costo_agua + costo_hielo + c_mixers_total + costo_fruta + C_EXTRA_BARRA

        ratio_barman = 40 if (checks['coctel_base'] or checks['coctel_prem']) else 50
        num_barmans = math.ceil(c.num_personas / ratio_barman)
        num_auxiliares = math.ceil(num_barmans / 2)
        if num_barmans > 1 and num_auxiliares == 0: num_auxiliares = 1

        C_BARMAN = self._get_costo(c.insumo_barman, 'COSTO_BARMAN', '1200.00')
        C_AUX = self._get_costo(c.insumo_auxiliar, 'COSTO_AUXILIAR', '800.00')
        costo_staff = (num_barmans * C_BARMAN) + (num_auxiliares * C_AUX)

        costo_total = res['costo_alcohol'] + res['costo_insumos_varios'] + costo_staff
        precio_sugerido = costo_total * Decimal(str(c.factor_utilidad_barra))

        costo_comun = costo_staff + costo_hielo + costo_agua + C_EXTRA_BARRA
        costo_puro['refrescos'] += c_mixers_total
        total_asignable = sum(costo_puro.values()) or 1
        desglose = {}
        margen = Decimal(str(c.factor_utilidad_barra))

        def get_linea(key):
            if costo_puro[key] > 0:
                participacion = costo_puro[key] / total_asignable
                full = costo_puro[key] + (costo_comun * participacion)
                return full * margen
            return Decimal(0)

        desglose['refrescos'] = get_linea('refrescos')
        desglose['cerveza'] = get_linea('cerveza')
        desglose['nacional'] = get_linea('nacional')
        desglose['premium'] = get_linea('premium')
        desglose['coctel'] = get_linea('coctel')

        diff = precio_sugerido - sum(desglose.values())
        if abs(diff) > 0.1:
            k = max(desglose, key=desglose.get)
            desglose[k] += diff

        return {
            'costo_total_estimado': costo_total,
            'precio_venta_sugerido_total': precio_sugerido,
            'botellas': res['botellas_nacional'] + res['botellas_premium'],
            'botellas_nacional': res['botellas_nacional'],
            'botellas_premium': res['botellas_premium'],
            'cervezas_unidades': res['cervezas_unidades'],
            'bolsas_hielo_20kg': res['bolsas_hielo_20kg'],
            'hielo_info': f"{int(res['hielo_consumo_kg'])}kg Consumo + {int(res['hielo_enfriamiento_kg'])}kg Frío",
            'litros_mezcladores': res['litros_mezcladores'],
            'litros_agua': litros_agua,
            'num_barmans': num_barmans,
            'num_auxiliares': num_auxiliares,
            'costo_alcohol': res['costo_alcohol'],
            'costo_insumos_varios': res['costo_insumos_varios'],
            'costo_hielo': costo_hielo,
            'costo_mixers_agua': c_mixers_total + costo_agua,
            'costo_fruta': costo_fruta,
            'costo_extra': C_EXTRA_BARRA,
            'costo_staff': costo_staff,
            'margen_aplicado': c.factor_utilidad_barra,
            'desglose_venta': desglose
        }


def actualizar_item_cotizacion(cotizacion):
    calc = CalculadoraBarraService(cotizacion)
    datos = calc.calcular()

    desc_clave = "Servicio de Barra"
    item_barra = cotizacion.items.filter(descripcion__startswith=desc_clave).first()

    if datos:
        precio = datos['precio_venta_sugerido_total']
        partes = []
        if cotizacion.incluye_cerveza: partes.append("Cerveza")
        if cotizacion.incluye_licor_nacional: partes.append("Nacional")
        if cotizacion.incluye_licor_premium: partes.append("Premium")
        if cotizacion.incluye_cocteleria_basica: partes.append("Cocteles")
        if cotizacion.incluye_cocteleria_premium: partes.append("Mixología")

        info = "/".join(partes) if partes else "Básico"
        clima_tag = "🔥" if cotizacion.clima in ['calor', 'extremo'] else ""
        nueva_desc = f"{desc_clave} [{info}] {clima_tag} | {cotizacion.num_personas} Pax - {cotizacion.horas_servicio} Hrs"

        if item_barra:
            if abs(item_barra.precio_unitario - precio) > Decimal('0.50') or item_barra.descripcion != nueva_desc:
                item_barra.precio_unitario = precio
                item_barra.descripcion = nueva_desc
                item_barra.cantidad = 1
                item_barra.save()
        else:
            ItemCotizacion.objects.create(
                cotizacion=cotizacion,
                descripcion=nueva_desc,
                cantidad=1,
                precio_unitario=precio
            )
    else:
        if item_barra: item_barra.delete()


# ==========================================
# PLAN DE PAGOS
# ==========================================
from datetime import timedelta
from django.utils import timezone


class PlanPagosService:
    """
    Genera planes de pago calendarizados según la anticipación del evento.
    """

    ESQUEMAS = {
        'largo':   {'min_dias': 120, 'parcialidades': [30, 25, 25, 20],
                    'conceptos': ['Anticipo', '2da Parcialidad', '3ra Parcialidad', 'Liquidación']},
        'medio':   {'min_dias': 60,  'parcialidades': [30, 35, 35],
                    'conceptos': ['Anticipo', '2da Parcialidad', 'Liquidación']},
        'corto':   {'min_dias': 30,  'parcialidades': [50, 50],
                    'conceptos': ['Anticipo (50%)', 'Liquidación']},
        'urgente': {'min_dias': 0,   'parcialidades': [50, 50],
                    'conceptos': ['Anticipo (50%)', 'Liquidación']},
    }

    DIAS_ANTES_ULTIMO_PAGO = 15

    def __init__(self, cotizacion):
        self.cotizacion = cotizacion

    def _get_esquema(self, dias_anticipacion):
        if dias_anticipacion >= 120: return self.ESQUEMAS['largo']
        elif dias_anticipacion >= 60: return self.ESQUEMAS['medio']
        elif dias_anticipacion >= 30: return self.ESQUEMAS['corto']
        else: return self.ESQUEMAS['urgente']

    def _generar_esquema_personalizado(self, num_parcialidades):
        porcentaje_base = round(100 / num_parcialidades, 2)
        porcentajes = [porcentaje_base] * num_parcialidades
        porcentajes[-1] = round(100 - sum(porcentajes[:-1]), 2)
        conceptos = []
        for i in range(num_parcialidades):
            if i == 0: conceptos.append('Anticipo')
            elif i == num_parcialidades - 1: conceptos.append('Liquidación')
            else: conceptos.append(f'Parcialidad {i + 1}')
        return {'parcialidades': porcentajes, 'conceptos': conceptos}

    def _calcular_fechas(self, fecha_contratacion, fecha_evento, num_parcialidades):
        fecha_ultimo_pago = fecha_evento - timedelta(days=self.DIAS_ANTES_ULTIMO_PAGO)
        hoy = timezone.now().date()
        if fecha_ultimo_pago <= hoy:
            fecha_ultimo_pago = fecha_evento - timedelta(days=3)
        if fecha_ultimo_pago <= hoy:
            fecha_ultimo_pago = hoy
        if num_parcialidades == 1:
            return [fecha_contratacion]
        if num_parcialidades == 2:
            return [fecha_contratacion, fecha_ultimo_pago]
        dias_total = (fecha_ultimo_pago - fecha_contratacion).days
        intervalo = dias_total / (num_parcialidades - 1)
        fechas = [fecha_contratacion]
        for i in range(1, num_parcialidades - 1):
            fechas.append(fecha_contratacion + timedelta(days=int(intervalo * i)))
        fechas.append(fecha_ultimo_pago)
        return fechas

    def generar(self, usuario=None, num_parcialidades=None):
        from .models import PlanPago, ParcialidadPago

        cotizacion = self.cotizacion
        monto_total = cotizacion.precio_final

        if monto_total <= 0:
            raise ValueError("La cotización no tiene precio final calculado.")

        PlanPago.objects.filter(cotizacion=cotizacion, activo=True).update(activo=False)

        hoy = timezone.now().date()
        dias_anticipacion = (cotizacion.fecha_evento - hoy).days

        if num_parcialidades and num_parcialidades >= 1:
            esquema = self._generar_esquema_personalizado(num_parcialidades)
        else:
            esquema = self._get_esquema(dias_anticipacion)

        porcentajes = esquema['parcialidades']
        conceptos = esquema['conceptos']
        fechas = self._calcular_fechas(hoy, cotizacion.fecha_evento, len(porcentajes))

        nota = f"Plan generado automáticamente. Anticipación: {dias_anticipacion} días."
        if num_parcialidades:
            nota = f"Plan personalizado de {num_parcialidades} parcialidades. Anticipación: {dias_anticipacion} días."

        plan = PlanPago.objects.create(
            cotizacion=cotizacion, generado_por=usuario, notas=nota
        )

        monto_acumulado = Decimal('0.00')
        for i, (porcentaje, concepto, fecha) in enumerate(zip(porcentajes, conceptos, fechas), 1):
            if i == len(porcentajes):
                monto = monto_total - monto_acumulado
            else:
                monto = (monto_total * Decimal(str(porcentaje)) / Decimal('100')).quantize(Decimal('0.01'))
                monto_acumulado += monto
            ParcialidadPago.objects.create(
                plan=plan, numero=i, concepto=concepto,
                monto=monto, porcentaje=Decimal(str(porcentaje)), fecha_limite=fecha,
            )

        return plan


# ==========================================
# GENERADOR DE CONTRATO DOCX
# ==========================================
class ContratoService:
    """
    Genera el contrato .docx a partir de una Cotizacion confirmada.
    """

    VERDE = "2E7D32"
    NEGRO = "1A1A1A"

    def __init__(self, cotizacion, tipo_servicio='EVENTO', deposito=Decimal('0.00')):
        self.cot  = cotizacion
        self.cli  = cotizacion.cliente
        self.tipo = tipo_servicio
        self.dep  = deposito

    def _fmt_fecha(self, d):
        if not d:
            return "—"
        meses = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio",
                 "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
        return f"{d.day} de {meses[d.month]} de {d.year}"

    def _fmt_hora(self, t):
        return t.strftime("%H:%M h") if t else "—"

    def _fmt_money(self, v):
        return f"${v:,.2f} MXN"

    def _tipo_servicio_display(self):
        return {"EVENTO": "Evento", "PASADIA": "Pasadía", "HOSPEDAJE": "Hospedaje"}.get(self.tipo, self.tipo)

    def _servicios_incluidos(self):
        partes = []
        for item in self.cot.items.select_related('producto').all():
            partes.append(item.descripcion or (item.producto.nombre if item.producto else ""))
        barra = []
        if self.cot.incluye_refrescos:          barra.append("Refrescos/Mezcladores")
        if self.cot.incluye_cerveza:            barra.append("Cerveza")
        if self.cot.incluye_licor_nacional:     barra.append("Licor Nacional")
        if self.cot.incluye_licor_premium:      barra.append("Licor Premium")
        if self.cot.incluye_cocteleria_basica:  barra.append("Coctelería")
        if self.cot.incluye_cocteleria_premium: barra.append("Mixología")
        if barra:
            partes.append("Barra: " + ", ".join(barra))
        return " | ".join(filter(None, partes)) or "Según cotización adjunta"

    def _primer_pago(self):
        return self.cot.pagos.order_by('fecha_pago').first()

    def _numero_contrato(self):
        from .models import ContratoServicio
        year = self.cot.fecha_evento.year if self.cot.fecha_evento else timezone.now().year
        seq  = ContratoServicio.objects.filter(
            cotizacion__fecha_evento__year=year
        ).count() + 1
        return f"CONT-{year}-{seq:04d}"

    def _add_bullet(self, doc, text):
        """Agrega un párrafo con viñeta manual, sin depender del estilo 'List Bullet'."""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = __import__('docx').shared.Cm(1)
        p.paragraph_format.first_line_indent = __import__('docx').shared.Cm(-0.5)
        run = p.add_run(f"• {text}")
        run.font.size = __import__('docx').shared.Pt(10)
        return p

    def generar(self):
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        import io

        doc = Document()

        # Márgenes del documento
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        # Estilo normal base
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(10)

        rv, gv, bv = int(self.VERDE[0:2], 16), int(self.VERDE[2:4], 16), int(self.VERDE[4:6], 16)
        color_verde = RGBColor(rv, gv, bv)

        # ── helpers locales ──────────────────────────────────────────────
        def add_heading(text, level=2):
            p = doc.add_paragraph()
            run = p.add_run(text.upper())
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11) if level == 2 else Pt(14)
            run.font.color.rgb = color_verde
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
            # Línea inferior simulada con borde de párrafo
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), self.VERDE)
            pBdr.append(bottom)
            pPr.append(pBdr)
            return p

        def add_field_table(rows_data):
            table = doc.add_table(rows=len(rows_data), cols=2)
            table.style = 'Table Grid'
            for i, (label, value) in enumerate(rows_data):
                lc = table.cell(i, 0)
                vc = table.cell(i, 1)
                lc.text = label
                vc.text = str(value)
                lr = lc.paragraphs[0].runs[0]
                lr.bold = True
                lr.font.size = Pt(9)
                vr = vc.paragraphs[0].runs[0]
                vr.font.size = Pt(9)
            for row in table.rows:
                row.cells[0].width = Cm(5)
                row.cells[1].width = Cm(11)
            return table

        def add_spacer(n=1):
            for _ in range(n):
                p = doc.add_paragraph("")
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)

        def add_bullet(text):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent       = Cm(1)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            r = p.add_run(f"\u2022  {text}")
            r.font.size = Pt(10)
            r.font.name = 'Arial'
            return p

        def add_body(text):
            p = doc.add_paragraph(text)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            if p.runs:
                p.runs[0].font.size = Pt(10)
            return p

        # ── ENCABEZADO ───────────────────────────────────────────────────
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = titulo.add_run("CONTRATO DE PRESTACIÓN DE SERVICIOS")
        r.bold = True
        r.font.size = Pt(16)
        r.font.name = 'Arial'
        r.font.color.rgb = color_verde

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rs = sub.add_run("Quinta Ko'ox Tanil  •  Yecapixtla, Morelos")
        rs.font.size = Pt(9)

        numero    = self._numero_contrato()
        folio     = f"COT-{self.cot.id:03d}"
        fecha_doc = self._fmt_fecha(timezone.now().date())

        ref = doc.add_paragraph()
        ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = ref.add_run(f"N° Contrato: {numero}  |  Folio ERP: {folio}  |  Fecha: {fecha_doc}")
        rr.font.size = Pt(8)

        add_spacer()

        # ── I. PARTES ─────────────────────────────────────────────────────
        add_heading("I. Partes")
        p = doc.add_paragraph()
        p.add_run("PRESTADOR DE SERVICIOS: ").bold = True
        p.add_run("Quinta Ko'ox Tanil, con domicilio en Yecapixtla, Morelos (en adelante \"LA QUINTA\").")

        p2 = doc.add_paragraph()
        p2.add_run("CONTRATANTE: ").bold = True
        p2.add_run("La persona cuyos datos se indican a continuación (en adelante \"EL CONTRATANTE\"):")

        add_spacer()
        nombre_fiscal = self.cli.razon_social or self.cli.nombre
        rfc_curp      = self.cli.rfc or "No proporcionado"
        add_field_table([
            ("Nombre / Razón Social", nombre_fiscal),
            ("RFC / CURP",            rfc_curp),
            ("Teléfono",              self.cli.telefono or "—"),
            ("Correo electrónico",    self.cli.email or "—"),
        ])
        add_spacer()

        # ── II. OBJETO ───────────────────────────────────────────────────
        add_heading("II. Objeto del Contrato")
        tipo_marca = {
            'EVENTO':    '☑ Evento     ☐ Pasadía     ☐ Hospedaje',
            'PASADIA':   '☐ Evento     ☑ Pasadía     ☐ Hospedaje',
            'HOSPEDAJE': '☐ Evento     ☐ Pasadía     ☑ Hospedaje',
        }.get(self.tipo, '☐ Evento     ☐ Pasadía     ☐ Hospedaje')

        hora_ini = self._fmt_hora(self.cot.hora_inicio)
        hora_fin = self._fmt_hora(self.cot.hora_fin)

        add_field_table([
            ("Tipo de servicio",    tipo_marca),
            ("Nombre del evento",   self.cot.nombre_evento),
            ("Fecha del evento",    self._fmt_fecha(self.cot.fecha_evento)),
            ("Hora de inicio",      hora_ini),
            ("Hora de término",     hora_fin),
            ("N° de personas",      str(self.cot.num_personas)),
            ("Servicios incluidos", self._servicios_incluidos()),
        ])
        add_spacer()

        # ── III. PRECIO Y PAGOS ──────────────────────────────────────────
        add_heading("III. Precio y Forma de Pago")
        primer_pago  = self._primer_pago()
        total_pagado = self.cot.total_pagado()
        saldo        = self.cot.saldo_pendiente()

        if primer_pago:
            metodos_map = dict(self.cot.pagos.model.METODOS)
            pago_str = (f"{self._fmt_money(primer_pago.monto)} — "
                        f"{metodos_map.get(primer_pago.metodo, primer_pago.metodo)} — "
                        f"{primer_pago.fecha_pago.strftime('%d/%m/%Y')}")
        else:
            pago_str = "Pendiente"

        add_field_table([
            ("Total del servicio",      self._fmt_money(self.cot.precio_final)),
            ("Anticipo pagado",         pago_str),
            ("Total pagado a la fecha", self._fmt_money(total_pagado)),
            ("Saldo pendiente",         self._fmt_money(saldo)),
            ("Depósito en garantía",    self._fmt_money(self.dep)),
            ("Datos bancarios",         "CLABE: ______________________________  Banco: __________"),
        ])
        add_spacer()

        p_av = doc.add_paragraph()
        p_av.paragraph_format.left_indent = Cm(0.3)
        r1 = p_av.add_run("⚠ IMPORTANTE: ")
        r1.bold = True
        r1.font.size = Pt(9)
        r2 = p_av.add_run("La fecha queda reservada ÚNICAMENTE al recibirse el anticipo. "
                           "Sin pago confirmado, LA QUINTA puede ceder la fecha a terceros.")
        r2.font.size = Pt(9)
        add_spacer()

        # ── IV. CANCELACIONES ────────────────────────────────────────────
        add_heading("IV. Política de Cancelación")
        cancel_data = [
            ("Más de 60 días",           "100% anticipo", "100% saldo"),
            ("30 a 59 días",             "50% anticipo",  "100% saldo"),
            ("15 a 29 días",             "0% anticipo",   "80% saldo"),
            ("Menos de 15 días",         "0% anticipo",   "0% saldo"),
            ("Fuerza mayor documentada", "100% anticipo", "100% saldo"),
        ]
        ct = doc.add_table(rows=len(cancel_data) + 1, cols=3)
        ct.style = 'Table Grid'
        for i, h in enumerate(["Tiempo antes del evento", "Reembolso anticipo", "Reembolso saldo"]):
            cell = ct.cell(0, i)
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        for ri, row_data in enumerate(cancel_data, 1):
            for ci, val in enumerate(row_data):
                c = ct.cell(ri, ci)
                c.text = val
                c.paragraphs[0].runs[0].font.size = Pt(9)
        add_spacer()

        # ── V. OBLIGACIONES ──────────────────────────────────────────────
        add_heading("V. Obligaciones de las Partes")
        p = doc.add_paragraph()
        p.add_run("5.1  LA QUINTA se obliga a:").bold = True
        for txt in [
            "Entregar el espacio limpio en el horario pactado.",
            "Proporcionar los servicios descritos en la Cláusula II.",
            "Designar un coordinador disponible durante el servicio.",
        ]:
            add_bullet(txt)

        add_spacer()
        p = doc.add_paragraph()
        p.add_run("5.2  EL CONTRATANTE se obliga a:").bold = True
        for txt in [
            "Pagar en los términos y plazos pactados.",
            "Respetar el horario acordado (horas extra con cargo previo acuerdo).",
            f"No exceder el aforo de {self.cot.num_personas} personas.",
            "Hacerse responsable del comportamiento de sus invitados y proveedores externos.",
            "No instalar infraestructura sin autorización escrita.",
            "Respetar el reglamento interno del lugar.",
        ]:
            add_bullet(txt)
        add_spacer()

        # ── VI. DAÑOS Y DEPÓSITO ─────────────────────────────────────────
        add_heading("VI. Daños y Depósito en Garantía")
        add_field_table([
            ("Depósito en garantía", self._fmt_money(self.dep)),
            ("Devolución",           "Dentro de 5 días hábiles post-evento, previo inventario."),
        ])
        add_spacer()

        # ── VII. DISPOSICIONES POR TIPO ──────────────────────────────────
        add_heading(f"VII. Disposiciones — {self._tipo_servicio_display()}")
        clausulas = {
            'EVENTO': [
                f"Música y sonido deben concluir a las {hora_fin}.",
                "Prohibido el acceso con pirotecnia, armas o sustancias ilícitas.",
                "Catering externo debe retirar todos sus residuos al concluir.",
            ],
            'PASADIA': [
                f"Acceso de {hora_ini} a {hora_fin}. No se permite pernoctar.",
                "Personas adicionales al aforo pactado tienen costo extra.",
            ],
            'HOSPEDAJE': [
                "Check-in y check-out en los horarios acordados.",
                "Prohibido organizar eventos sin contrato independiente.",
                "Fumar en habitaciones genera cargo de limpieza extraordinaria.",
            ],
        }
        for txt in clausulas.get(self.tipo, []):
            add_bullet(txt)
        add_spacer()

        # ── VIII. JURISDICCIÓN ───────────────────────────────────────────
        add_heading("VIII. Jurisdicción")
        p = doc.add_paragraph()
        r = p.add_run(
            "Las partes se someten a los Tribunales del Estado de Morelos, renunciando a cualquier "
            "otro fuero. Ante controversia intentarán resolución amigable en 10 días hábiles antes "
            "de acción legal. Adicionalmente podrán acudir a PROFECO como instancia de mediación."
        )
        r.font.size = Pt(9)
        add_spacer()

        # ── IX. ACEPTACIÓN Y FIRMAS ──────────────────────────────────────
        add_heading("IX. Aceptación y Firmas")
        p = doc.add_paragraph()
        r = p.add_run(
            "Leído y entendido en todas sus partes, ambas partes lo aceptan y se obligan a su "
            "cumplimiento. La firma puede ser física o electrónica (WhatsApp/correo con valor probatorio)."
        )
        r.font.size = Pt(9)
        add_spacer(2)

        firma_table = doc.add_table(rows=3, cols=3)
        firma_table.cell(0, 0).text = "_" * 30
        firma_table.cell(0, 2).text = "_" * 30
        firma_table.cell(1, 0).text = nombre_fiscal
        firma_table.cell(1, 2).text = "Quinta Ko'ox Tanil"
        firma_table.cell(2, 0).text = "CONTRATANTE"
        firma_table.cell(2, 2).text = "PRESTADOR DE SERVICIOS"
        for ri in range(3):
            for ci in [0, 2]:
                cell_p = firma_table.cell(ri, ci).paragraphs[0]
                cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if cell_p.runs:
                    cell_p.runs[0].bold = (ri > 0)
                    cell_p.runs[0].font.size = Pt(9)

        add_spacer()
        pf = doc.add_paragraph(f"Fecha: {fecha_doc}")
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ── Serializar ───────────────────────────────────────────────────
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read(), numero